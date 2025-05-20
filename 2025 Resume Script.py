import json
import os
import winsound
from datetime import datetime
from docx import Document

def load_json("C:\Users\irato\OneDrive\Documents\company.json"
              "C:\Users\irato\OneDrive\Documents\with application.py"):
    try:
        with open(file_path, 'r') as file:
            return json.load(file)
    except (FileNotFoundError, json.JSONDecodeError) as e:
        print(f"Error loading {file_path}: {e}")
        return {}

def get_user_input():
    file_path = r"C:\Users\irato\OneDrive\Documents\company.json"
    if os.path.exists(file_path):
        return load_json(file_path)
    else:
        print("File not found. Please check the path and try again.")
        return {}
        return {}

# Prompt user for job description and resume
print("Please provide the job description.")
job_description = get_user_input()

print("Please provide the resume.")
resume = load_json(r"C:\Users\irato\OneDrive\Documents\with application.py")

# Check if job description and resume are loaded correctly
if not job_description or not resume:
    print("Error: Job description or resume data is missing or invalid.")
    exit(1)

# Extract key skills from job description
required_skills = job_description.get('skills', [])

# Match skills with resume
matched_skills = [skill for skill in resume.get('skills', []) if skill in required_skills]

# Generate tailored resume with expanded job duties
tailored_resume = {
    "name": resume['name'],
    "contact": resume['contact'],
    "skills": matched_skills,
    "experience": [
        {
            "company": "ABC Corp",
            "role": "Data Analyst",
            "years": 2,
            "duties": [
                "Analyzed large datasets to identify trends and insights, improving decision-making processes.",
                "Developed and maintained dashboards and reports using Python and SQL.",
                "Collaborated with cross-functional teams to define data requirements and deliver actionable insights.",
                "Conducted data cleaning and preprocessing to ensure data accuracy and integrity."
            ]
        },
        {
            "company": "XYZ Ltd",
            "role": "Project Manager",
            "years": 3,
            "duties": [
                "Led project teams to successfully deliver projects on time and within budget.",
                "Managed project timelines, resources, and budgets to ensure project goals were met.",
                "Facilitated communication between stakeholders and project teams to ensure alignment and transparency.",
                "Implemented project management best practices to improve efficiency and effectiveness."
            ]
        }
    ],
    "education": resume['education']
}

# Generate tailored cover letter
cover_letter = f"""
Dear {job_description['company']},

I am writing to express my enthusiasm for the Software Engineer Apprenticeship position at {job_description['company']}. With a strong background in customer service and a passion for technology, I am excited about the opportunity to contribute to your team.

In my previous role as a Customer Service Representative at XYZ Corp, I honed my problem-solving and communication skills by assisting customers with technical issues and ensuring their satisfaction. My ability to troubleshoot and resolve complex problems has been a key factor in improving customer satisfaction scores by 20% over the past year.

To transition into the IT field, I have completed several courses in Python, Machine Learning, and Data Analysis. These courses have equipped me with the technical skills necessary to excel in this role. Additionally, my experience in customer service has provided me with a unique perspective on user needs and the importance of creating user-friendly solutions.

I am particularly drawn to {job_description['company']} because of your commitment to innovation and excellence. I am confident that my skills in {', '.join(matched_skills)}, combined with my dedication to continuous learning, will enable me to make a meaningful contribution to your team.

I have attached my resume for your review. Thank you for considering my application. I look forward to the opportunity to discuss how my background and skills align with the goals of {job_description['company']}.

Sincerely,
{resume['name']}
"""

# Generate timestamp for filenames
timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")

# Define output directory
output_dir = r"C:\Users\irato\OneDrive - Minnesota State\Documents\Resumes & Cover Letters"

# Ensure output directory exists
os.makedirs(output_dir, exist_ok=True)

# Define filenames with timestamps
resume_filename = os.path.join(output_dir, f"python_resume_{timestamp}.docx")
cover_letter_filename = os.path.join(output_dir, f"cover_letter_{timestamp}.docx")

# Save tailored resume in DOCX format
resume_doc = Document()
resume_doc.add_heading('Resume', 0)
resume_doc.add_heading('Name', level=1)
resume_doc.add_paragraph(tailored_resume['name'])
resume_doc.add_heading('Contact', level=1)
resume_doc.add_paragraph(tailored_resume['contact'])
resume_doc.add_heading('Skills', level=1)
resume_doc.add_paragraph(', '.join(tailored_resume['skills']))
resume_doc.add_heading('Experience', level=1)
for exp in tailored_resume['experience']:
    resume_doc.add_heading(exp['company'], level=2)
    resume_doc.add_paragraph(f"Role: {exp['role']}")
    resume_doc.add_paragraph(f"Years: {exp['years']}")
    resume_doc.add_heading('Duties', level=3)
    for duty in exp['duties']:
        resume_doc.add_paragraph(duty, style='List Bullet')
resume_doc.add_heading('Education', level=1)
resume_doc.add_paragraph(tailored_resume['education'])
resume_doc.save(resume_filename)

# Save tailored cover letter in DOCX format
cover_letter_doc = Document()
cover_letter_doc.add_paragraph(cover_letter)
cover_letter_doc.save(cover_letter_filename)

print(f"Tailored resume saved as: {resume_filename}")
print(f"Cover letter saved as: {cover_letter_filename}")

# Play a chime sound
winsound.MessageBeep()